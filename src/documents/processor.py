"""
Document processing functionality for extracting and chunking content
"""
import hashlib
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime

import PyPDF2
import docx
from bs4 import BeautifulSoup

from ..config import get_logger, settings

logger = get_logger(__name__)


class DocumentProcessor:
    """Processes uploaded documents to extract text and create chunks"""
    
    def __init__(self):
        self.supported_types = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.doc': self._extract_docx,
            '.txt': self._extract_text,
            '.md': self._extract_text,
            '.html': self._extract_html,
            '.htm': self._extract_html
        }
        
        self.chunk_size = 1000  # Characters per chunk
        self.chunk_overlap = 200  # Overlap between chunks
    
    async def process_document(self, file_path: Path, metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Process a document file and extract text chunks
        
        Args:
            file_path: Path to the document file
            metadata: Optional metadata about the document
            
        Returns:
            Dict containing extracted text, chunks, and metadata
        """
        try:
            logger.info("Processing document", file_path=str(file_path))
            
            # Get file information
            file_info = self._get_file_info(file_path)
            
            # Extract text based on file type
            text_content = await self._extract_text_content(file_path)
            
            if not text_content:
                raise ValueError(f"No text content extracted from {file_path}")
            
            # Create chunks
            chunks = self._create_chunks(text_content)
            
            # Generate document ID
            doc_id = self._generate_document_id(file_path, text_content)
            
            # Prepare metadata
            doc_metadata = {
                "id": doc_id,
                "filename": file_info["filename"],
                "file_type": file_info["file_type"],
                "file_size": file_info["file_size"],
                "processed_at": datetime.utcnow().isoformat(),
                "chunk_count": len(chunks),
                "total_chars": len(text_content),
                "title": metadata.get("title", file_info["filename"]) if metadata else file_info["filename"]
            }
            
            if metadata:
                doc_metadata.update(metadata)
            
            result = {
                "document_id": doc_id,
                "metadata": doc_metadata,
                "full_text": text_content,
                "chunks": chunks,
                "success": True
            }
            
            logger.info("Document processed successfully",
                       document_id=doc_id,
                       chunks=len(chunks),
                       chars=len(text_content))
            
            return result
            
        except Exception as e:
            logger.error("Error processing document", 
                        file_path=str(file_path),
                        error=str(e))
            return {
                "document_id": None,
                "metadata": {},
                "full_text": "",
                "chunks": [],
                "success": False,
                "error": str(e)
            }
    
    async def process_url_content(self, url: str, content: str, title: str = None) -> Dict[str, Any]:
        """
        Process content from a URL
        
        Args:
            url: The source URL
            content: The extracted content
            title: Optional title for the content
            
        Returns:
            Dict containing processed content and metadata
        """
        try:
            logger.info("Processing URL content", url=url)
            
            # Clean the content
            text_content = self._clean_text(content)
            
            if not text_content:
                raise ValueError(f"No text content from URL {url}")
            
            # Create chunks
            chunks = self._create_chunks(text_content)
            
            # Generate document ID
            doc_id = self._generate_url_id(url, text_content)
            
            # Prepare metadata
            doc_metadata = {
                "id": doc_id,
                "source_url": url,
                "file_type": "url",
                "processed_at": datetime.utcnow().isoformat(),
                "chunk_count": len(chunks),
                "total_chars": len(text_content),
                "title": title or f"Web content from {url}"
            }
            
            result = {
                "document_id": doc_id,
                "metadata": doc_metadata,
                "full_text": text_content,
                "chunks": chunks,
                "success": True
            }
            
            logger.info("URL content processed successfully",
                       document_id=doc_id,
                       chunks=len(chunks),
                       chars=len(text_content))
            
            return result
            
        except Exception as e:
            logger.error("Error processing URL content", 
                        url=url,
                        error=str(e))
            return {
                "document_id": None,
                "metadata": {},
                "full_text": "",
                "chunks": [],
                "success": False,
                "error": str(e)
            }
    
    def _get_file_info(self, file_path: Path) -> Dict[str, Any]:
        """Get basic file information"""
        stat = file_path.stat()
        
        # Guess MIME type
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        return {
            "filename": file_path.name,
            "file_size": stat.st_size,
            "file_type": file_path.suffix.lower(),
            "mime_type": mime_type,
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat()
        }
    
    async def _extract_text_content(self, file_path: Path) -> str:
        """Extract text content from file based on type"""
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        extractor = self.supported_types[file_extension]
        text_content = await extractor(file_path)
        
        return self._clean_text(text_content)
    
    async def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        text_parts = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(text)
                    except Exception as e:
                        logger.warning("Error extracting page from PDF", 
                                     page=page_num,
                                     error=str(e))
                        continue
        
        except Exception as e:
            logger.error("Error reading PDF", file_path=str(file_path), error=str(e))
            raise
        
        return '\n\n'.join(text_parts)
    
    async def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error("Error reading DOCX", file_path=str(file_path), error=str(e))
            raise
    
    async def _extract_text(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                return file.read()
                
        except Exception as e:
            logger.error("Error reading text file", file_path=str(file_path), error=str(e))
            raise
    
    async def _extract_html(self, file_path: Path) -> str:
        """Extract text from HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            return text
            
        except Exception as e:
            logger.error("Error reading HTML file", file_path=str(file_path), error=str(e))
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            cleaned_line = ' '.join(line.split())  # Normalize whitespace
            if cleaned_line:  # Only keep non-empty lines
                cleaned_lines.append(cleaned_line)
        
        # Join with single newlines
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove excessive newlines (more than 2)
        import re
        cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
        
        return cleaned_text.strip()
    
    def _create_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Create overlapping chunks from text content"""
        if not text:
            return []
        
        chunks = []
        start = 0
        chunk_id = 0
        
        while start < len(text):
            # Calculate end position
            end = start + self.chunk_size
            
            # If this would be the last chunk, take all remaining text
            if end >= len(text):
                chunk_text = text[start:]
            else:
                # Try to break at word boundary
                chunk_text = text[start:end]
                last_space = chunk_text.rfind(' ')
                if last_space > self.chunk_size * 0.8:  # Only break at word if reasonable
                    chunk_text = chunk_text[:last_space]
                    end = start + last_space
            
            if chunk_text.strip():  # Only add non-empty chunks
                chunk = {
                    "id": chunk_id,
                    "text": chunk_text.strip(),
                    "start_char": start,
                    "end_char": start + len(chunk_text),
                    "char_count": len(chunk_text)
                }
                chunks.append(chunk)
                chunk_id += 1
            
            # Move start position with overlap
            if end >= len(text):
                break
            
            start = max(start + 1, end - self.chunk_overlap)
        
        logger.debug("Created text chunks", 
                    total_chunks=len(chunks),
                    avg_size=sum(c["char_count"] for c in chunks) // len(chunks) if chunks else 0)
        
        return chunks
    
    def _generate_document_id(self, file_path: Path, content: str) -> str:
        """Generate a unique document ID based on file and content"""
        # Combine file path, size, and content hash
        file_info = f"{file_path.name}_{file_path.stat().st_size}"
        content_hash = hashlib.md5(content.encode('utf-8')).hexdigest()[:8]
        
        return f"doc_{file_info}_{content_hash}"
    
    def _generate_url_id(self, url: str, content: str) -> str:
        """Generate a unique document ID for URL content"""
        url_hash = hashlib.md5(url.encode('utf-8')).hexdigest()[:8]
        content_hash = hashlib.md5(content.encode('utf-8')).hexdigest()[:8]
        
        return f"url_{url_hash}_{content_hash}"
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        return list(self.supported_types.keys())